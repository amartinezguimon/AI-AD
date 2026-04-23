"""
build_docx.py  —  converts docs/executive_report.md into docs/executive_report.docx
with all images embedded, tables formatted, headings styled and references intact.

Run from the project root:
    python docs/build_docx.py
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPORT_MD  = "docs/executive_report.md"
FIGURES_DIR = "docs/figures"
OUTPUT_DOCX = "docs/executive_report.docx"


# ── helpers ──────────────────────────────────────────────────────────────────

def set_heading_style(para, level):
    """Apply a built-in heading style by level."""
    style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}
    para.style = style_map.get(level, "Heading 3")


def add_horizontal_rule(doc):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_break()


def set_cell_bg(cell, hex_color):
    """Set table cell background colour (hex without #)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def apply_bold_inline(para, text):
    """
    Parse **bold** and `code` inline markers and add runs with correct formatting.
    Everything else is plain.
    """
    # combined pattern: **bold**, `code`, or plain text
    pattern = re.compile(r"(\*\*.*?\*\*|`[^`]+`)")
    parts = pattern.split(text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        else:
            if part:
                para.add_run(part)


def add_paragraph_with_inline(doc, text, style="Normal"):
    # strip leading/trailing * used for italic wrapping in captions
    text = text.strip("*").strip()
    para = doc.add_paragraph(style=style)
    apply_bold_inline(para, text)
    return para


def add_image(doc, img_path, caption_text):
    """Add an image centred with a caption below."""
    if not os.path.exists(img_path):
        doc.add_paragraph(f"[Image not found: {img_path}]").italic = True
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(img_path, width=Inches(5.5))

    cap = doc.add_paragraph(caption_text.strip("*").strip())
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.style = "Caption" if "Caption" in [s.name for s in doc.styles] else "Normal"
    for run in cap.runs:
        run.italic = True
        run.font.size = Pt(9)
    doc.add_paragraph()   # breathing room


def add_md_table(doc, rows):
    """
    rows: list of lists of strings.
    First row is the header. Second row is the separator (--|--|--), skip it.
    """
    data_rows = [r for r in rows if not re.match(r"^\s*[-|: ]+\s*$", r[0])]
    if not data_rows:
        return

    col_count = max(len(r) for r in data_rows)
    table = doc.add_table(rows=len(data_rows), cols=col_count)
    table.style = "Table Grid"

    for i, row in enumerate(data_rows):
        for j, cell_text in enumerate(row):
            if j >= col_count:
                break
            cell = table.cell(i, j)
            cell_text = cell_text.strip()
            # strip bold markers for cell display
            cell_text = re.sub(r"\*\*(.*?)\*\*", r"\1", cell_text)
            cell_text = re.sub(r"`([^`]+)`", r"\1", cell_text)
            cell.text = cell_text
            if i == 0:
                set_cell_bg(cell, "D9E2F3")
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True

    doc.add_paragraph()


def parse_md_table_row(line):
    """Split a markdown table row into cells."""
    line = line.strip().strip("|")
    return [c.strip() for c in line.split("|")]


# ── main converter ────────────────────────────────────────────────────────────

def build_docx():
    with open(REPORT_MD, encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()

    # page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # default body font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    n = len(lines)

    # accumulate table rows
    table_buffer = []
    in_table     = False

    def flush_table():
        nonlocal table_buffer, in_table
        if table_buffer:
            add_md_table(doc, table_buffer)
        table_buffer = []
        in_table     = False

    while i < n:
        raw  = lines[i]
        line = raw.rstrip("\n")

        # ── horizontal rule ──────────────────────────────────────────────────
        if line.strip() == "---":
            flush_table()
            add_horizontal_rule(doc)
            i += 1
            continue

        # ── headings ─────────────────────────────────────────────────────────
        m = re.match(r"^(#{1,3})\s+(.*)", line)
        if m:
            flush_table()
            level = len(m.group(1))
            text  = re.sub(r"\*\*(.*?)\*\*", r"\1", m.group(2))  # strip bold
            para  = doc.add_heading(text, level=level)
            i += 1
            continue

        # ── images ───────────────────────────────────────────────────────────
        m = re.match(r"^!\[(.*?)\]\((.*?)\)", line)
        if m:
            flush_table()
            alt      = m.group(1)
            rel_path = m.group(2)
            img_path = os.path.join("docs", rel_path)
            # look ahead for caption (italic line immediately after)
            caption = alt
            if i + 1 < n:
                next_line = lines[i + 1].strip()
                if next_line.startswith("*") and next_line.endswith("*"):
                    caption = next_line.strip("*").strip()
                    i += 1   # consume caption line
            add_image(doc, img_path, caption)
            i += 1
            continue

        # ── markdown tables ───────────────────────────────────────────────────
        if "|" in line and line.strip().startswith("|"):
            cells = parse_md_table_row(line)
            table_buffer.append(cells)
            in_table = True
            i += 1
            continue
        else:
            if in_table:
                flush_table()

        # ── blank line ────────────────────────────────────────────────────────
        if line.strip() == "":
            i += 1
            continue

        # ── bullet list ───────────────────────────────────────────────────────
        m = re.match(r"^(\s*)[*\-]\s+(.*)", line)
        if m:
            indent = len(m.group(1))
            text   = m.group(2)
            style_name = "List Bullet 2" if indent > 0 else "List Bullet"
            para = doc.add_paragraph(style=style_name)
            apply_bold_inline(para, text)
            i += 1
            continue

        # ── numbered list ─────────────────────────────────────────────────────
        m = re.match(r"^\d+\.\s+\*\*(.*?)\.\*\*\s*(.*)", line)
        if m:
            label = m.group(1)
            rest  = m.group(2)
            para  = doc.add_paragraph(style="List Number")
            bold_run = para.add_run(label + ". ")
            bold_run.bold = True
            apply_bold_inline(para, rest)
            i += 1
            continue

        m = re.match(r"^\d+\.\s+(.*)", line)
        if m:
            para = doc.add_paragraph(style="List Number")
            apply_bold_inline(para, m.group(1))
            i += 1
            continue

        # ── bold-only lines (used as sub-labels) ──────────────────────────────
        if line.strip().startswith("**") and line.strip().endswith("**"):
            para = doc.add_paragraph()
            run  = para.add_run(line.strip().strip("*"))
            run.bold = True
            i += 1
            continue

        # ── caption lines (italic, starts with *Figure or *Table) ─────────────
        if line.strip().startswith("*Figure") or line.strip().startswith("*Table"):
            cap = doc.add_paragraph(line.strip().strip("*").strip())
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cap.runs:
                run.italic = True
                run.font.size = Pt(9)
            i += 1
            continue

        # ── reference lines (<a id=...>) — skip the anchor tag ────────────────
        if line.strip().startswith("<a id="):
            # strip html anchor, keep text after it
            text = re.sub(r"<a[^>]+></a>", "", line).strip()
            if text:
                para = doc.add_paragraph(style="Normal")
                apply_bold_inline(para, text)
            i += 1
            continue

        # ── default: normal paragraph ─────────────────────────────────────────
        # handle leading ">" blockquote style
        if line.strip().startswith(">"):
            line = line.strip().lstrip("> ")

        para = doc.add_paragraph(style="Normal")
        apply_bold_inline(para, line.strip())
        i += 1

    flush_table()

    os.makedirs(os.path.dirname(OUTPUT_DOCX), exist_ok=True)
    doc.save(OUTPUT_DOCX)
    print(f"Saved: {OUTPUT_DOCX}")


if __name__ == "__main__":
    build_docx()
